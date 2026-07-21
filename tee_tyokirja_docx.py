#!/usr/bin/env python3
"""Generoi kurssin ladattava Word-työkirja.

Markdown-versio downloads/ai-perusteet-tyokirja.md on sisällöllinen pari.
Aja repon juuressa:

    python3 tee_tyokirja_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "downloads" / "ai-perusteet-tyokirja.docx"

# compact_reference_guide-pohja, nimetty A4- ja kurssibrändin mukainen
# workbook_form-override.
PAGE_W_MM = 210
PAGE_H_MM = 297
MARGIN_MM = 18
HEADER_MM = 9
FOOTER_MM = 10
CONTENT_W_DXA = round((PAGE_W_MM - 2 * MARGIN_MM) / 25.4 * 1440)

FONT_BODY = "Arial"
FONT_HEAD = "Georgia"
FONT_META = "Arial"

DARK = "1B2336"
BODY = "3A4253"
MUTED = "5A6478"
HERO = "0B0F1A"
WHITE = "FFFFFF"
KICKER = "7E88A8"
BLUE = "3B5BDB"
VIOLET = "863FC4"
TEAL = "008799"
LINE = "C9D0DF"
BORDER = "D8DEEC"
PAPER = "F4F6FB"
BLUE_TINT = "EEF1FE"
VIOLET_TINT = "F3EDFA"
TEAL_TINT = "E9F6F7"
GOLD_TINT = "FFF8E8"
GOLD = "8A6A00"

MODULES = {
    "theory": ("OSA 1 · TEORIA", BLUE, BLUE_TINT),
    "use": ("OSA 2 · TEKOÄLYJEN KÄYTTÖ", VIOLET, VIOLET_TINT),
    "agents": ("OSA 3 · AGENTIT", TEAL, TEAL_TINT),
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name=FONT_BODY, size=10.5, color=BODY, *, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def set_cell_margins(cell, *, top=120, bottom=120, start=210, end=210):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    # Word käyttää start/end-arvoja, mutta kaikki katselu- ja muunnosohjelmat
    # eivät tunnista niitä. Kirjoitetaan rinnalle left/right-varmistukset, jotta
    # tekstin ja pystyviivan väli säilyy eri ohjelmissa.
    for side, value in (
        ("top", top),
        ("bottom", bottom),
        ("start", start),
        ("end", end),
        ("left", start),
        ("right", end),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths, *, indent=210, margins=None):
    widths = [int(x) for x in widths]
    if sum(widths) != CONTENT_W_DXA:
        widths[-1] += CONTENT_W_DXA - sum(widths)
    margins = margins or {"top": 120, "bottom": 120, "start": 210, "end": 210}
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_W_DXA))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            set_cell_margins(cell, **margins)


def set_table_borders(table, color=BORDER, size=6, *, inside=True, left_color=None, left_size=None):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        visible = inside or edge not in ("insideH", "insideV")
        node.set(qn("w:val"), "single" if visible else "nil")
        node.set(qn("w:sz"), str(left_size if edge == "left" and left_size else size))
        node.set(qn("w:color"), left_color if edge == "left" and left_color else color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color, size=24, space=12):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, FONT_META, 8, KICKER)


def configure_section(section, right_label="TYÖKIRJA", *, first_page=False):
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.header_distance = Mm(HEADER_MM)
    section.footer_distance = Mm(FOOTER_MM)
    section.different_first_page_header_footer = first_page

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(PAGE_W_MM - 2 * MARGIN_MM))
    set_run_font(p.add_run("AI · PERUSTEET | TYÖKIRJA"), FONT_META, 8, KICKER, bold=True)
    p.add_run("\t")
    set_run_font(p.add_run(right_label), FONT_META, 8, KICKER, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(PAGE_W_MM - 2 * MARGIN_MM))
    set_run_font(p.add_run("aiperusteet.fi · CC BY-SA 4.0"), FONT_META, 8, KICKER)
    p.add_run("\t")
    add_page_number(p)


def add_body(doc, text, *, size=10.5, color=BODY, bold=False, italic=False, after=5, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = keep
    set_run_font(p.add_run(text), FONT_BODY, size, color, bold=bold, italic=italic)
    return p


def add_kicker(doc, text, color=KICKER, *, after=4, page_break_before=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    set_run_font(p.add_run(text.upper()), FONT_META, 8.5, color, bold=True)
    return p


def add_h1(doc, text, color=BLUE, *, after=8):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    # Borderin w:space ei toteudu yhdenmukaisesti kaikissa Word-lukijoissa.
    # Kiinteä sarkain tekee pystyviivan ja otsikkotekstin välistä luotettavan.
    p.paragraph_format.tab_stops.add_tab_stop(Mm(4.5))
    paragraph_left_border(p, color, 28, 0)
    p.add_run("\t")
    set_run_font(p.add_run(text), FONT_HEAD, 22, DARK, bold=True)
    return p


def add_h2(doc, text, color=BLUE, *, before=12, after=5):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), FONT_HEAD, 15.5, color, bold=True)
    return p


def add_h3(doc, text, color=DARK, *, before=8, after=3):
    # Korttien ja taulukoiden väliotsikot ovat asiakirjarakenteessa H2-tasoa,
    # vaikka ne ladotaan varsinaisia osio-otsikoita pienemmällä kirjasimella.
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), FONT_BODY, 11, color, bold=True)
    return p


def add_page_heading(doc, module_key, title, intro=None, *, sublabel=None, page_break_before=False):
    label, color, _ = MODULES[module_key]
    add_kicker(doc, sublabel or label, color, page_break_before=page_break_before)
    add_h1(doc, title, color)
    if intro:
        add_body(doc, intro, after=7)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_info_box(doc, title, text, *, color=BLUE, fill=PAPER):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [CONTENT_W_DXA])
    set_table_borders(table, BORDER, 5, inside=False, left_color=color, left_size=24)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title), FONT_BODY, 10, DARK, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text), FONT_BODY, 9.5, BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_field(doc, label, *, hint=None, height=18, color=BLUE, placeholder="Kirjoita tähän."):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(label), FONT_BODY, 10.5, DARK, bold=True)
    if hint:
        set_run_font(p.add_run(f"  {hint}"), FONT_BODY, 8.5, MUTED, italic=True)
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [CONTENT_W_DXA])
    set_table_borders(table, LINE, 5, inside=False, left_color=color, left_size=10)
    cell = table.cell(0, 0)
    shade_cell(cell, "FBFCFE")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    row = table.rows[0]
    row.height = Mm(height)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(placeholder), FONT_BODY, 9, KICKER, italic=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_checklist(doc, items, *, color=BLUE, fill=PAPER):
    table = doc.add_table(rows=len(items), cols=1)
    apply_table_geometry(table, [CONTENT_W_DXA])
    set_table_borders(table, BORDER, 4, inside=True, left_color=color, left_size=12)
    for row, item in zip(table.rows, items):
        row.height = Mm(9)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cell = row.cells[0]
        shade_cell(cell, fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("☐  "), FONT_BODY, 11, color, bold=True)
        set_run_font(p.add_run(item), FONT_BODY, 9.5, BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_choice_box(doc, title, items, *, color=BLUE, fill=PAPER):
    add_h3(doc, title, color)
    add_checklist(doc, items, color=color, fill=fill)


def add_prompt_card(doc, *, color=VIOLET):
    add_kicker(doc, "RAKENNUSPALIKKA 1", color)
    add_h1(doc, "Testattu promptikortti", color)
    add_body(doc, "Laadi yksi uudelleen käytettävä promptikortti. Testaa kaksi versiota samalla aineistolla ja muuta niiden välissä vain yhtä nimettyä asiaa.")
    add_field(doc, "Kortin nimi ja käyttötilanne", height=12, color=color)
    add_field(doc, "Tavoite, käyttäjä tai yleisö sekä vaihtuvan syötteen paikka", height=20, color=color)
    add_field(doc, "Laatukriteerit, jotka päätit ennen ensimmäistä ajoa", height=20, color=color)
    add_field(doc, "Promptin versio 1, käytetty aineisto ja saatu vastaus", height=34, color=color)

    page_break(doc)
    add_page_heading(doc, "use", "Promptikortin vertailu", "Dokumentoi yksi perusteltu muutos ja rajaa johtopäätös tehtyyn kokeeseen.", sublabel="RAKENNUSPALIKKA 1 · JATKUU")
    add_field(doc, "Nimeä yksi havaittu puute ja sitä vastaava promptimuutos.", height=18, color=color)
    add_field(doc, "Promptin versio 2 ja samalla aineistolla saatu vastaus", height=34, color=color)
    add_field(doc, "Mitä muutos osoitti tässä testissä? Kirjaa myös kortin tunnettu raja ja yhteys botin järjestelmäpromptiin.", height=26, color=color)


def add_label_detail_table(doc, rows, *, color=BLUE, fill="FBFCFE", label_width=2350, row_height=10):
    table = doc.add_table(rows=len(rows), cols=2)
    widths = [label_width, CONTENT_W_DXA - label_width]
    apply_table_geometry(table, widths)
    set_table_borders(table, BORDER, 5, inside=True, left_color=color, left_size=10)
    for row, (label, value) in zip(table.rows, rows):
        row.height = Mm(row_height)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], WHITE)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), FONT_BODY, 9, DARK, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), FONT_BODY, 9, KICKER, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_test_card(doc, number, category, *, color=TEAL):
    add_h3(doc, f"Testi {number} · {category}", color, before=5, after=3)
    add_label_detail_table(
        doc,
        [
            ("Nimi", "Kirjoita tähän."),
            ("Syöte", "Kirjoita tähän."),
            ("Odotettu tulos", "Kirjoita tähän."),
            ("Todellinen tulos", "Kirjoita tähän."),
            ("Tila ja huomiot", "Läpäisi / ei läpäissyt — perustele lyhyesti."),
        ],
        color=color,
        row_height=7.5,
    )


def add_source_card(doc, number, *, color=VIOLET):
    add_h3(doc, f"Tietopohjan dokumentti {number}", color, before=5, after=3)
    add_label_detail_table(
        doc,
        [
            ("Dokumentti ja linkki", "Kirjoita tähän."),
            ("Miksi mukana?", "Kirjoita tähän."),
            ("Tietotarve", "Kirjoita tähän."),
            ("Luotettavuus", "Arvioi ajantasaisuus ja lähteen luotettavuus."),
            ("Oikeudet", "Kirjaa tietosuoja ja käyttöoikeus."),
        ],
        color=color,
        row_height=8,
    )


def add_node_table(doc):
    table = doc.add_table(rows=6, cols=3)
    widths = [1500, 4100, CONTENT_W_DXA - 5600]
    apply_table_geometry(table, widths)
    set_table_borders(table, BORDER, 5, inside=True, left_color=TEAL, left_size=10)
    headers = ["Solmu", "Tehtävä, syöte ja tulos", "Yhteys rakennusosiin"]
    for i, text in enumerate(headers):
        shade_cell(table.rows[0].cells[i], TEAL_TINT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), FONT_BODY, 8.5, DARK, bold=True)
    for row in table.rows[1:]:
        row.height = Mm(17)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run("Kirjoita tähän."), FONT_BODY, 8.5, KICKER, italic=True)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def new_module_section(doc, module_key):
    label, _, _ = MODULES[module_key]
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, label)
    return section


def build_document():
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BODY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = FONT_HEAD if style_name != "Heading 3" else FONT_BODY
        styles[style_name].font.color.rgb = rgb(DARK)

    configure_section(doc.sections[0], "TYÖKIRJA", first_page=True)

    # 1 · Kansi ja aloitus
    hero = doc.add_paragraph()
    hero.paragraph_format.space_before = Pt(0)
    hero.paragraph_format.space_after = Pt(8)
    hero.paragraph_format.left_indent = Mm(5)
    hero.paragraph_format.right_indent = Mm(5)
    shade_paragraph(hero, HERO)
    set_run_font(hero.add_run("AI · PERUSTEET — AVOIN VERKKOKURSSI"), FONT_META, 9, WHITE, bold=True)
    add_h1(doc, "Työkirja", BLUE, after=2)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(14)
    set_run_font(title.add_run("Kokeet, havainnot ja lopputyöt"), FONT_HEAD, 16, MUTED, italic=True)
    add_body(doc, "Työkirja kulkee mukanasi koko kurssin ajan. Kirjaa promptit, kokeiden tulokset, lähteet ja tekemäsi korjaukset työkirjaan työskentelyn aikana. Kurssisivusto ei säilytä näitä tuotoksia puolestasi.", size=11, after=7)
    add_body(doc, "Työkirja ei korvaa oppituntien tehtävänantoja. Avaa aina myös kyseisen tunnin ohje. Tämä pohja auttaa sinua kokoamaan arvioinnissa tarvittavan näytön yhteen paikkaan. Voit täyttää koko työkirjan tai vain opiskelemasi osion.", after=9)
    add_info_box(doc, "Huolehdi tietosuojasta", "Älä kirjoita työkirjaan henkilötietoja, salaisuuksia tai organisaation luottamuksellista tietoa. Käytä harjoituksissa kuvitteellista tai anonymisoitua aineistoa.", color=GOLD, fill=GOLD_TINT)
    add_h2(doc, "Omat tiedot ja opiskelutapa", BLUE, before=9)
    add_label_detail_table(
        doc,
        [
            ("Nimi tai sovittu tunniste", "Kirjoita tähän."),
            ("Opiskelutapa", "Itsenäisesti / oppilaitoksen toteutuksessa"),
            ("Täytän osat", "Teoria / Tekoälyjen käyttö / Agentit"),
            ("Tallennuspaikka", "Kirjoita tähän."),
            ("Käyttämäni työkalut", "Kirjoita tähän. Työkalu voi vaihtua tehtävästä toiseen."),
        ],
        color=BLUE,
        row_height=10,
    )
    add_info_box(doc, "Kolme osaa — kolme mahdollista projektia", "Sinun ei tarvitse käyttää samaa aihetta kaikissa osissa. Valitse kuhunkin lopputyöhön tilanne, jonka avulla voit osoittaa kyseisen osion tavoitteiden mukaisen osaamisen.", color=BLUE, fill=BLUE_TINT)

    page_break(doc)
    add_kicker(doc, "TYÖKIRJAN KARTTA", BLUE)
    add_h1(doc, "Näin etenet itsenäisesti", BLUE)
    add_body(doc, "Työkirja on yksi kokonaisuus, mutta sinun ei tarvitse täyttää sitä kerralla. Valitse opiskelemasi osa ja etene oppituntien tahdissa.", after=8)
    add_info_box(doc, "Teoria", "Kerää todistusaineistot tunneilla 3, 5 ja 7. Viimeistele lopputyö tunnilla 9.", color=BLUE, fill=BLUE_TINT)
    add_info_box(doc, "Tekoälyjen käyttö", "Rakenna promptauspankki sekä botin määrittely ja tietopohja tunneilla 12, 14 ja 15. Viimeistele apuri-botti tunneilla 17–18.", color=VIOLET, fill=VIOLET_TINT)
    add_info_box(doc, "Agentit", "Tee pohjapiirrokset tunneilla 19, 21 ja 23–25. Rakenna, testaa ja viimeistele kokonaisuus tunneilla 26–27.", color=TEAL, fill=TEAL_TINT)
    add_h2(doc, "Työskentele neljässä vaiheessa", BLUE, before=7)
    add_checklist(doc, ["Avaa aina myös kyseisen oppitunnin tehtävänanto kurssisivustolta.", "Täytä vain opiskelemasi osan ajankohtainen kohta. Lisää kuvat, kaaviot ja laajat tuotokset linkkeinä tai liitteinä.", "Tallenna työ säännöllisesti omalle laitteellesi tai valitsemaasi pilvipalveluun.", "Tee osion lopputarkistus ja vie valmis työ myös PDF-muotoon."], color=BLUE, fill=PAPER)
    add_info_box(doc, "Liiku Wordissa otsikoiden avulla", "Valitse Näytä → Siirtymisruutu. Otsikkoluettelosta pääset suoraan haluamaasi moduuliin ja tehtävään.", color=BLUE, fill=BLUE_TINT)

    # OSA 1 · TEORIA
    new_module_section(doc, "theory")
    add_page_heading(doc, "theory", "Todistusaineisto 1", "Tee kokeet tunnilla 3. Säilytä tarkat syötteet ja tulokset, jotta voit palata niihin Teoria-osion lopputyössä.", sublabel="OSA 1 · TEORIA · TUNTI 3")
    add_h2(doc, "Koe A: Seuraavan sanan ennustaminen", BLUE)
    add_field(doc, "Kirjoita tarkka syöte.", height=22, color=BLUE)
    for n in range(1, 4):
        add_field(doc, f"Tulos {n}", height=18, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Sama kysymys, eri vastaus", "Toista sama kysymys ja vertaa vastauksia keskenään.", sublabel="TODISTUSAINEISTO 1 · JATKUU")
    add_h2(doc, "Koe B: Sama kysymys, eri vastaus", BLUE, before=10)
    add_field(doc, "Kirjoita tarkka kysymys.", height=18, color=BLUE)
    add_field(doc, "Vastaus 1", height=22, color=BLUE)
    add_field(doc, "Vastaus 2", height=22, color=BLUE)
    add_field(doc, "Vertaa vastauksia. Mitä eroja huomasit?", height=24, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Vakuuttava vai oikea?", "Tarkista mallin vastaus luotettavasta lähteestä ja tee oma analyysisi.", sublabel="TODISTUSAINEISTO 1 · JATKUU")
    add_field(doc, "Kirjoita tarkistettava kysymys ja mallin vastaus.", height=38, color=BLUE)
    add_field(doc, "Nimeä luotettava lähde ja kirjoita tarkistettu vastaus.", height=32, color=BLUE)
    add_field(doc, "Mitä havainto kertoo kielimallin luotettavuudesta?", height=28, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Oma analyysi", "Kokoa kokeiden havainnot yhdeksi perustelluksi selitykseksi.", sublabel="TODISTUSAINEISTO 1 · JATKUU")
    add_h2(doc, "Oma analyysi, noin 300 sanaa", BLUE)
    add_body(doc, "Selitä, miten kielimalli tuottaa tekstiä. Kerro, mitä tarkoittaa, että malli ennustaa eikä tiedä, ja kuvaa tilanne, jossa tästä voi syntyä ongelma. Käytä käsitteitä token, seuraavan tokenin ennustaminen, lämpötila tai epädeterminismi ja hallusinaatio.")
    add_field(doc, "Kirjoita analyysi.", height=68, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Todistusaineisto 2", "Tee koe tunnilla 5. Osoita, miten rajaus voi kadota kontekstista ja miten käyttäjä palauttaa sen.", sublabel="OSA 1 · TEORIA · TUNTI 5")
    add_field(doc, "Kirjoita alkuperäinen rajaus kokonaisuudessaan.", height=32, color=BLUE)
    add_field(doc, "Millä kysymyksellä testasit rajauksen muistamista?", height=22, color=BLUE)
    add_field(doc, "Kirjoita vastaus ennen rajauksen palauttamista.", height=30, color=BLUE)
    add_field(doc, "Kuvaa, milloin ja miten rajaus katosi.", height=25, color=BLUE)
    add_field(doc, "Kirjoita toistettu rajaus ja sen jälkeinen vastaus.", height=38, color=BLUE)

    add_page_heading(doc, "theory", "Oma tapausesimerkki", "Tarkastele kontekstin rajoja yhden käytännön tilanteen kautta.", sublabel="TODISTUSAINEISTO 2 · JATKUU", page_break_before=True)
    add_h2(doc, "Oma tapausesimerkki, noin 300 sanaa", BLUE)
    add_body(doc, "Kuvaa tilanne, jossa olennainen tieto unohtui. Selitä, miksi näin kävi, mitä siitä seurasi ja miten käyttäjä voi ehkäistä ongelman. Käytä käsitteitä konteksti-ikkuna, FIFO-periaate, ankkurointi, pilkkominen ja dokumentointi.")
    add_field(doc, "Kirjoita tapausesimerkki.", height=55, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Todistusaineisto 3", "Arvioi tunnilla 7 sekä kielimallin väite että petoksia tunnistavan luokittelumallin tulokset.", sublabel="OSA 1 · TEORIA · TUNTI 7")
    add_h2(doc, "Kielimallin väitteen tarkistus", BLUE, before=7)
    add_field(doc, "Kirjaa tarkistettava väite ja mallin vastaus.", height=34, color=BLUE)
    add_field(doc, "Nimeä riippumaton lähde ja kirjaa, mitä se osoittaa.", height=30, color=BLUE)
    add_field(doc, "Tee johtopäätös: vahvistuiko väite, kumoutuiko se vai jäikö se epävarmaksi?", height=30, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Luokittelumallin arvio", "Petosmalli arvioi 1 000 maksua. Todellisia petoksia oli 10. Malli löysi niistä 8 ja pysäytti lisäksi 42 tavallista maksua.", sublabel="TODISTUSAINEISTO 3 · JATKUU")
    add_field(doc, "Kirjaa väärät positiiviset ja väärät negatiiviset sekä selitä niiden käytännön seuraukset.", height=44, color=BLUE)
    add_field(doc, "Laske tai tulkitse kokonaistarkkuus, osumatarkkuus ja kattavuus. Mitä kukin mittari kertoo — ja mitä se ei kerro?", height=64, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Oma tarkistuskäytäntö", "Kirjoita 6–8 vaihetta, joilla tarkistat tekoälytuloksen. Aloita jokainen vaihe toimintaverbillä.", sublabel="TODISTUSAINEISTO 3 · JATKUU")
    add_field(doc, "Kirjoita 6–8 vaihetta.", height=115, color=BLUE)
    add_info_box(doc, "Pidä tarkistuskäytäntö käytännöllisenä", "Esimerkiksi: tunnista tehtävä ja virhetyyppi, lukitse odotus, etsi alkuperäinen lähde, vertaa tietoja ja mittareita, arvioi seuraus, merkitse epävarmuus ja kirjaa ihmisen tarkistus.", color=BLUE, fill=BLUE_TINT)

    page_break(doc)
    add_page_heading(doc, "theory", "Teoria-osion lopputyö", "Valitse kirjallinen tai visuaalis-suullinen suoritustapa. Käytä kaikkia kolmea todistusaineistoa.", sublabel="OSA 1 · TEORIA · LOPPUTYÖ")
    add_choice_box(doc, "Valitse suoritustapa", ["Kirjallinen asiantuntijalausunto, 500–700 sanaa", "Tilannekartta ja päätöspuu sekä 5–8 minuutin suullinen esitys"], color=BLUE, fill=BLUE_TINT)
    add_choice_box(doc, "Valitse skenaario", ["Asiakaspalvelu ja tietovuoto", "Rekrytointialgoritmi ja syrjintä", "Markkinointisisältö ja tekijänoikeudet"], color=BLUE, fill=BLUE_TINT)
    add_info_box(doc, "Rajaa lähdetyö", "Käytä tunnilla 8 arvioimiasi lähteitä. Jos jokin työn keskeinen väite ei selviä niistä, tee tunnilla 9 enintään yksi uusi lähdetarkistus.", color=BLUE, fill=BLUE_TINT)
    add_field(doc, "Kirjaa keskeiset riskit ja käsitteet.", height=32, color=BLUE)
    add_field(doc, "Nimeä käyttämäsi todistusaineistot ja tunnilla 8 arvioidut lähteet. Kirjaa, mitä väitettä kukin lähde tukee.", height=30, color=BLUE)
    add_field(doc, "Liitä tarkka prompti tai keskusteluloki.", height=24, color=BLUE, placeholder="Kirjoita tähän tai lisää linkki tai liite.")

    add_page_heading(doc, "theory", "Työn neljä osaa", "Jäsennä valitsemasi suoritustapa näiden neljän kysymyksen avulla.", sublabel="TEORIA-OSION LOPPUTYÖ · JATKUU", page_break_before=True)
    add_field(doc, "1. Mitä tapahtui? Kuvaa tilanne neutraalisti.", height=24, color=BLUE)
    add_field(doc, "2. Miksi se tapahtui? Selitä syyt kurssin käsitteillä.", height=28, color=BLUE)
    add_field(doc, "3. Miten tilanne pitäisi hoitaa? Ehdota konkreettiset toimenpiteet.", height=30, color=BLUE)
    add_field(doc, "4. Miten ammatillinen ja käytännön vastuu jakautuvat? Kerro, kuka vastaa mistä ja miksi.", height=28, color=BLUE)

    page_break(doc)
    add_page_heading(doc, "theory", "Palautettava työ", "Täytä valitsemaasi suoritustapaan kuuluvat kohdat ja säilytä myös työprosessisi näyttö.", sublabel="TEORIA-OSION LOPPUTYÖ · TUOTOS")
    add_h2(doc, "Kirjallinen suoritustapa", BLUE)
    add_field(doc, "Lisää luonnos ja valmis asiantuntijalausunto linkkeinä tai liitteinä.", height=22, color=BLUE, placeholder="Lisää linkit tai liitteet.")
    add_h2(doc, "Visuaalis-suullinen suoritustapa", BLUE)
    add_info_box(doc, "Kevyt tukirunko", "Jaa yhden sivun tilannekartta neljään kenttään: mitä tapahtui, miksi se tapahtui, miten tilanne hoidetaan ja miten vastuu jakautuu. Tee päätöspuuhun kolme päätöskohtaa, kaksi haaraa kuhunkin sekä päätepisteisiin toiminta ja vastuurooli. Avainsanat riittävät, koska avaat perustelut puheessa.", color=BLUE, fill=BLUE_TINT)
    add_field(doc, "Tallenna yhden sivun tilannekartta.", height=18, color=BLUE, placeholder="Lisää linkki tai liite.")
    add_field(doc, "Tallenna kevyt päätöspuu.", height=18, color=BLUE, placeholder="Lisää linkki tai liite.")
    add_field(doc, "Tallenna 5–8 minuutin esityksen muistiinpanot tai nauhoite.", height=24, color=BLUE, placeholder="Kirjoita tähän tai lisää linkki tai liite.")
    add_info_box(doc, "Muista prosessinäyttö", "Säilytä lähteet, luonnos ja tekemäsi korjaukset. Valmis työ ei yksin osoita, miten tarkistit väitteet ja päädyit lopputulokseen.", color=BLUE, fill=BLUE_TINT)

    page_break(doc)
    add_page_heading(doc, "theory", "Korjausloki ja puolustus", "Dokumentoi kaksi ennen–jälkeen-korjausta. Perustele kumpikin korjaus lähteellä tai kurssin käsitteellä.", sublabel="TEORIA-OSION LOPPUTYÖ · VIIMEISTELY")
    for n in range(1, 3):
        add_h3(doc, f"Korjaus {n}", BLUE)
        add_label_detail_table(doc, [("Ennen", "Kirjoita tähän."), ("Jälkeen", "Kirjoita tähän."), ("Perustelu", "Miksi korjasit? Nimeä lähde tai käsite.")], color=BLUE, row_height=7)
    add_field(doc, "Kirjoita kirjallisen suoritustavan 1–2 minuutin puolustuksen muistiinpanot.", height=26, color=BLUE)
    add_page_heading(doc, "theory", "Lopputarkistus", "Varmista ennen palautusta, että työ osoittaa sekä ymmärryksesi että työprosessisi.", sublabel="TEORIA-OSION LOPPUTYÖ · VIIMEISTELY", page_break_before=True)
    add_checklist(doc, ["Käytän käsitteitä täsmällisesti ja sovellan niitä tilanteeseen.", "Hyödynnän kaikkia kolmea todistusaineistoa.", "Ehdotan konkreettisia toimenpiteitä: kuka tekee, mitä tekee ja milloin.", "Tuotos on selkeä ja asiallinen, ja sen väitteet ovat perusteltuja.", "Käsittelen vastuukysymyksen omana osanaan."], color=BLUE, fill=BLUE_TINT)

    # OSA 2 · TEKOÄLYJEN KÄYTTÖ
    new_module_section(doc, "use")
    add_prompt_card(doc)

    page_break(doc)
    add_page_heading(doc, "use", "Botin määrittely", "Rakenna botin perustamisasiakirja tunnilla 14. Kirjoita selkeästi, kenelle botti on, mitä se tekee ja missä sen rajat kulkevat.", sublabel="OSA 2 · TEKOÄLYJEN KÄYTTÖ · TUNTI 14")
    add_field(doc, "Botin nimi", height=12, color=VIOLET)
    add_field(doc, "Botin aihe", height=14, color=VIOLET)
    add_field(doc, "Tyypillinen käyttötilanne: missä tilanteessa ja kenelle botti on hyödyllinen?", height=18, color=VIOLET)
    add_field(doc, "Kohderyhmä", height=18, color=VIOLET)
    add_field(doc, "Tarkoitus: mitä käyttäjän ongelmaa botti auttaa ratkaisemaan?", height=28, color=VIOLET)
    add_field(doc, "Persoona ja äänensävy", height=22, color=VIOLET)

    page_break(doc)
    add_page_heading(doc, "use", "Rajaukset ja palaute", "Tee botin rajat näkyviksi ja dokumentoi, miten palaute muutti suunnitelmaa.", sublabel="BOTIN MÄÄRITTELY · JATKUU")
    add_field(doc, "Kuvaa työnkulku 5–7 vaiheena.", height=38, color=VIOLET)
    add_field(doc, "Kirjaa 3–5 rajaa sille, mitä botti ei tee.", height=35, color=VIOLET)
    add_field(doc, "Mitä rakentavaa tai kriittistä palautetta sait tekoälyltä tai vertaiselta? Mitä muutit sen perusteella?", height=30, color=VIOLET)
    add_field(doc, "Tiivistä botin ydin yhteen lauseeseen.", height=16, color=VIOLET)

    page_break(doc)
    add_page_heading(doc, "use", "Tietopohja", "Valitse tunnilla 15 yhteensä 2–4 dokumenttia. Perustele niiden hyöty, luotettavuus, ajantasaisuus ja käyttöoikeus.", sublabel="OSA 2 · TEKOÄLYJEN KÄYTTÖ · TUNTI 15")
    add_field(doc, "Kirjaa 5–8 tietotarvetta, joihin botin pitää osata vastata.", height=38, color=VIOLET)
    for n in range(1, 3):
        add_source_card(doc, n)

    page_break(doc)
    add_page_heading(doc, "use", "Tietopohjan dokumentit", "Täydennä loput dokumentit ja arvioi tietopohjan kattavuus.", sublabel="RAKENNUSPALIKKA 3 · JATKUU")
    for n in range(3, 5):
        add_source_card(doc, n)
    add_page_heading(doc, "use", "Tietopohjan kattavuus", "Arvioi kokonaisuutta yksittäisten dokumenttien laadun lisäksi.", sublabel="RAKENNUSPALIKKA 3 · YHTEENVETO", page_break_before=True)
    add_field(doc, "Mitä tietopohja kattaa hyvin? Mitä se ei vielä kata?", height=28, color=VIOLET)
    add_info_box(doc, "Tunnista myös puuttuva tieto", "Hyvä tietopohja ei yritä näyttää kattavammalta kuin se on. Kirjaa näkyviin, milloin botin pitää sanoa, ettei tieto riitä vastaukseen.", color=VIOLET, fill=VIOLET_TINT)
    add_field(doc, "Tietopohjan kansio tai tallennuspaikka", height=18, color=VIOLET)
    add_checklist(doc, ["Tietopohjan 2–4 varsinaista lähdedokumenttia on tallennettu samaan kansioon, ja niiden nimet vastaavat taulukkoa."], color=VIOLET, fill=VIOLET_TINT)

    page_break(doc)
    add_page_heading(doc, "use", "Yhden välineen kokeilu", "Kokeile tunnilla 16 yhtä työkalureittiä käytännössä. Jos sopivaa palvelua ei ole, arvioi opettajan antamia esimerkkituotoksia.", sublabel="OSA 2 · TEKOÄLYJEN KÄYTTÖ · TUNTI 16")
    add_choice_box(doc, "Valitse kokeiltu reitti", ["Kuva", "Ääni tai musiikki", "Video", "Koodi"], color=VIOLET, fill=VIOLET_TINT)
    add_field(doc, "Käyttötarkoitus, käyttäjä ja tärkein onnistumisen ehto", height=20, color=VIOLET)
    add_field(doc, "Käytetty väline, syöte ja mahdollinen lähdeaineisto", height=28, color=VIOLET)
    add_field(doc, "Ennalta päätetyt arviointikriteerit ja version 1 havainnot", height=28, color=VIOLET)
    add_field(doc, "Yksi nimetty muutos, version 2 tulos ja rajattu johtopäätös", height=28, color=VIOLET)

    page_break(doc)
    add_page_heading(doc, "use", "Bottiprojektin valintakortti", "Vertaa työkalureittejä ja kahta toteutustapaa. Tee lopuksi toteutus-, riski- ja käyttöönottopäätös.", sublabel="TUNTI 16 · TALLENNETTAVA TUOTOS")
    add_label_detail_table(doc, [("Kuva", "Sopiiko käyttäjätarpeeseen? Miksi tai miksi ei?"), ("Ääni tai musiikki", "Sopiiko käyttäjätarpeeseen? Miksi tai miksi ei?"), ("Video", "Sopiiko käyttäjätarpeeseen? Miksi tai miksi ei?"), ("Koodi", "Sopiiko käyttäjätarpeeseen? Miksi tai miksi ei?")], color=VIOLET, row_height=9)
    add_field(doc, "Vertaa toteutustapoja A ja B: järjestelmäprompti, tietopohja, mahdollinen erikoistyökalu, näyttö sekä suurin riski.", height=38, color=VIOLET)
    add_field(doc, "Kirjoita toteutuspäätös, tärkein riskirajaus ja havaittava käyttöönottokriteeri. Nimeä suunnittelupolulla teknisesti todentamattomat ominaisuudet.", height=34, color=VIOLET)

    page_break(doc)
    add_page_heading(doc, "use", "Apuri-botin lopputyö", "Valitse toimiva tekninen toteutus tai dokumentoitu suunnittelusuoritus. Polut ovat samanarvoisia, mutta niissä osoitetaan osaaminen eri tavalla.", sublabel="OSA 2 · TEKOÄLYJEN KÄYTTÖ · LOPPUTYÖ")
    add_choice_box(doc, "Valitse toteutustapa", ["Toimiva tekninen toteutus", "Dokumentoitu suunnittelusuoritus"], color=VIOLET, fill=VIOLET_TINT)
    add_field(doc, "Lisää teknisen botin linkki ja suoritusjälki tai suunnittelusuorituksen arkkitehtuuri ja simuloitu suoritusjälki. Merkitse simuloidut vaiheet sekä teknisesti todentamattomat ominaisuudet.", height=24, color=VIOLET)
    add_field(doc, "Liitä päivitetty botin määrittely.", height=24, color=VIOLET, placeholder="Kirjoita tähän tai lisää linkki tai liite.")
    add_field(doc, "Kirjoita järjestelmäprompti kokonaisuudessaan.", height=70, color=VIOLET)

    page_break(doc)
    add_page_heading(doc, "use", "Tietopohjan perustelut", "Nimeä lopulliseen toteutukseen valitsemasi dokumentit ja perustele, miksi juuri ne kuuluvat mukaan.", sublabel="APURI-BOTIN LOPPUTYÖ · JATKUU")
    add_field(doc, "Nimeä tietopohjan 2–4 dokumenttia ja perustele valinnat.", height=70, color=VIOLET)
    add_info_box(doc, "Tarkista ennen testausta", "Varmista, että järjestelmäprompti, botin määrittely ja tietopohja ovat keskenään johdonmukaisia. Korjaa ristiriidat ennen testien ajamista.", color=VIOLET, fill=VIOLET_TINT)

    page_break(doc)
    add_page_heading(doc, "use", "Testaus, korjaus ja reflektio", "Aja normaali tapaus, kielteinen testi ja reunatapaus. Tee yksi nimetty korjaus ja aja juuri sitä koskeva testi uudelleen samalla odotuksella.", sublabel="APURI-BOTIN LOPPUTYÖ · VIIMEISTELY")
    for n, category in enumerate(("Normaali tapaus", "Kielteinen testi", "Reunatapaus"), 1):
        add_test_card(doc, n, category, color=VIOLET)
    add_h2(doc, "Nimetty korjaus ja uudelleentesti", VIOLET)
    add_label_detail_table(doc, [("Havaittu puute", "Kirjoita tähän."), ("Tehty muutos", "Kirjoita tähän."), ("Ennen ja jälkeen", "Vertaa tuloksia ja kirjoita johtopäätös.")], color=VIOLET, row_height=11)
    add_field(doc, "Reflektio, 200–300 sanaa tai saavutettava äänite/selostettu kuvakooste", hint="Mitä opit, mikä havainto johti korjaukseen ja mitä tekisit seuraavaksi?", height=45, color=VIOLET)
    add_field(doc, "Esittelyn muistiinpanot, 2–3 minuuttia", hint="Käyttäjä ja ongelma, yksi ennen–jälkeen-korjaus, yksi perusteltu rajaus sekä vastauksesi mahdolliseen jatkokysymykseen.", height=28, color=VIOLET)
    add_h2(doc, "Lopputarkistus", VIOLET)
    add_checklist(doc, ["Tekninen botti toimii rajatussa tehtävässä tai suunnittelusuoritus kuvaa toteuttamiskelpoisen rakenteen ja rajaa teknisesti todentamattomat ominaisuudet.", "Järjestelmäpromptissa on selkeä rooli, työnkulku ja rajat.", "Tietopohjassa on 2–4 perusteltua dokumenttia.", "Normaali, kielteinen ja reunatapaus sekä yksi korjaus ja sen uudelleentesti on dokumentoitu.", "Reflektio ja esittely osoittavat, että ymmärrän tekemäni valinnat."], color=VIOLET, fill=VIOLET_TINT)

    # OSA 3 · AGENTIT
    new_module_section(doc, "agents")
    add_page_heading(doc, "agents", "Agentin pohjapiirrokset", "Tällä kurssilla rakennettavalla tekoälyagentilla tarkoitetaan kielimallin ja sitä ohjaavan agentin ohjauskehyksen (harness) muodostamaa järjestelmää. Suunnittelun tarkistuslistaan kuuluu kuusi rakennusosaa.", sublabel="OSA 3 · AGENTIT · TUNNIT 19–25")
    add_info_box(doc, "Kuusi rakennusosaa", "Syötekäsittelijä · päättelijä ja suunnittelija · työkalujen suorittaja · muisti ja konteksti · turvakerros · seuranta ja palautesilmukka", color=TEAL, fill=TEAL_TINT)
    add_h2(doc, "Pohjapiirros 1: Ongelma", TEAL)
    add_body(doc, "Kirjoita 150–200 sanaa. Kuvaa ongelma ja käyttäjä, perustele, miksi ratkaisu tarvitsee agentin eikä tavallista automaatiota tai chatbotia, ja arvioi kuuden rakennusosan tarvetta. Perustele myös, jos jätät jonkin niistä pois.")
    add_field(doc, "Kirjoita pohjapiirros.", height=52, color=TEAL)
    add_h2(doc, "Pohjapiirros 2: Muisti", TEAL)
    add_body(doc, "Kirjoita 150–200 sanaa. Kuvaa yhden suorituksen konteksti, mahdollinen pitkäkestoinen muisti sekä prosessin tilat ja tilasiirtymät. Perustele, mitä tallennetaan, kuinka pitkäksi aikaa ja miksi. Älä sijoita agentin roolia tai toimintaohjeita muistiin.")
    add_field(doc, "Kirjoita pohjapiirros.", height=52, color=TEAL)
    add_h2(doc, "Järjestelmäohjeet — erillään muistista", TEAL)
    add_body(doc, "Kirjaa agentin tehtävä, rooli, sallitut toimintatavat ja rajat järjestelmäohjeisiin. Järjestelmäohjeet ohjaavat toimintaa, mutta eivät ole keskusteluhistoriaa, prosessin tilaa tai pitkäkestoista muistia.")
    add_field(doc, "Kirjoita järjestelmäohjeiden ydin.", height=34, color=TEAL)

    page_break(doc)
    add_page_heading(doc, "agents", "Päättely ja turva", "Kuvaa valintasi niin, että toinen ihminen pystyy ymmärtämään ja tarkistamaan työnkulun.", sublabel="AGENTIN POHJAPIIRROKSET · JATKUU")
    add_h2(doc, "Pohjapiirros 3: Päättely", TEAL)
    add_body(doc, "Kirjoita 150–200 sanaa. Valitse ReAct tai eksplisiittinen työnkulku ja perustele valintasi. Kuvaa, miten päättely näkyy työnkulun rakenteessa ja lokissa. Määritä kummassakin mallissa vähintään yksi aito rajattu kielimallivalinta vähintään kahdesta sallitusta vaihtoehdosta. Anna lokiesimerkki, jossa näkyvät syöte, sallitut vaihtoehdot, mallin valinta, tulos tai virhe ja seuraava toiminto. Älä tallenna kielimallin raakaa päättelyketjua.")
    add_field(doc, "Kirjoita pohjapiirros.", height=52, color=TEAL)
    add_h2(doc, "Pohjapiirros 4: Turva", TEAL)
    add_body(doc, "Kirjoita 150–200 sanaa. Määritä, mitä agentti saa ja ei saa tehdä. Tunnista 2–3 riskiä. Kuvaa, miten validoit ja rajaat toimintaa, seuraat tapahtumia, palaudut virheistä ja määrität lokiin tallennettavat tiedot.")
    add_field(doc, "Kirjoita pohjapiirros.", height=52, color=TEAL)

    page_break(doc)
    add_page_heading(doc, "agents", "Ihminen ja toteutus", "Arvioi hyväksynnän tarve ennen kuin rakennat työnkulun. Pidä agentti rajattuna ja käytä vain vastuiden kuvaamiseen tarvittavat vaiheet.", sublabel="AGENTIN POHJAPIIRROKSET · JATKUU")
    add_choice_box(doc, "Työskentelytapa", ["Teen työn yksin", "Teen työn parin kanssa"], color=TEAL, fill=TEAL_TINT)
    add_field(doc, "Parin nimi ja työnjako", height=20, color=TEAL, placeholder="Täytä, jos teet työn parin kanssa.")
    add_h2(doc, "Pohjapiirros 5: Ihminen", TEAL)
    add_body(doc, "Kirjoita 150–200 sanaa. Tunnista hyväksyntää vaativat kohdat. Jos niitä ei ole, perustele ratkaisu riskien perusteella. Jos hyväksyntäportteja on, kerro, mitä tietoja kukin portti näyttää ja missä kanavassa hyväksyntää pyydetään. Tietojen pitää riittää päätökseen noin 30 sekunnissa. Valitse jokaiselle portille varapolku siltä varalta, ettei hyväksyntäpyyntöön vastata: oletushylkäys, vain matalan riskin tilanteeseen sopiva oletushyväksyntä tai eskalointi toiselle ihmiselle.")
    add_field(doc, "Kirjoita pohjapiirros.", height=52, color=TEAL)

    page_break(doc)
    add_page_heading(doc, "agents", "Koottu suunnitelma", "Tarkista viisi pohjapiirrosta yhtenä kokonaisuutena ennen rakentamista.", sublabel="AGENTIN POHJAPIIRROKSET · JOHDONMUKAISUUS")
    add_checklist(doc, ["Ongelma ja työnkulku vastaavat toisiaan.", "Muistiratkaisu noudattaa tietosuojarajauksia.", "Päättelymalli ja lokitus tukevat toisiaan.", "Turvakerros vastaa työkalujen oikeuksia ja riskejä.", "Mahdolliset hyväksyntäportit ja varapolut ovat yksiselitteisiä."], color=TEAL, fill=TEAL_TINT)
    add_field(doc, "Mitkä valinnat toistuvat kaikissa viidessä pohjapiirroksessa?", height=26, color=TEAL)
    add_field(doc, "Missä havaitsit ristiriidan tai aukon?", height=28, color=TEAL)
    add_field(doc, "Mitä päivitit ennen rakentamista?", height=28, color=TEAL)

    page_break(doc)
    add_page_heading(doc, "agents", "Toteutus", "Pidä agentti yksinkertaisena. Piirrä pääpolku ensin ja kuvaa vasta sitten yksittäiset solmut.", sublabel="OSA 3 · AGENTIT · RAKENTAMINEN")
    add_field(doc, "Lisää työnkulun linkki, JSON-vientitiedosto tai alustariippumaton suoritusjälki.", height=18, color=TEAL)
    add_field(doc, "Piirrä rajattu työnkulku. Näytä mallin valinta, työkalut, turva, hyväksyntä, lokitus ja virhepolku siellä, missä ne vaikuttavat.", height=78, color=TEAL, placeholder="Lisää kaavio tähän.")
    add_page_heading(doc, "agents", "Solmuluettelo", "Kuvaa jokainen solmu niin, että työnkulku on ymmärrettävissä myös ilman toteutusalustaa.", sublabel="OSA 3 · AGENTIT · RAKENTAMINEN", page_break_before=True)
    add_node_table(doc)

    page_break(doc)
    add_page_heading(doc, "agents", "Näyttöpaketti", "Kokoa yhden 90 minuutin oppitunnin aikana tiivis näyttö. Tekninen ja dokumentoitu polku ovat samanarvoisia.", sublabel="OSA 3 · AGENTIT · LOPPUTYÖ")
    add_h2(doc, "Kuuden rakennusosan tarkistus", TEAL)
    add_body(doc, "Merkitse, missä syötekäsittelijä, päättelijä ja suunnittelija, työkalujen suorittaja, muisti ja konteksti, turvakerros sekä seuranta ja palautesilmukka näkyvät tai miksi niitä ei tarvita.")
    add_field(doc, "Kirjaa rakennusosien kattavuus.", height=48, color=TEAL)
    add_h2(doc, "Aito rajattu mallivalinta", TEAL)
    add_label_detail_table(doc, [("Syöte", "Kirjoita tähän."), ("Vähintään kaksi sallittua vaihtoehtoa", "Kirjoita tähän."), ("Mallin todellinen valinta", "Kirjoita tähän."), ("Seuraava vaihe", "Kirjoita tähän.")], color=TEAL, row_height=8)
    add_info_box(doc, "Polkujen näyttö", "Tekninen polku käyttää suoritusnäkymää. Dokumentoitu polku näyttää todellisen mallikutsun ja merkitsee sitä seuraavat simuloidut vaiheet.", color=TEAL, fill=TEAL_TINT)

    add_page_heading(doc, "agents", "Kolme testiä", "Aja yksi normaali, yksi reuna- ja yksi turvallisuustesti. Kirjaa syöte, odotettu tulos, todellinen tulos ja johtopäätös.", sublabel="OSA 3 · AGENTIT · TESTAUS", page_break_before=True)
    for number, category in enumerate(("Normaali", "Reunatapaus", "Turvallisuus"), 1):
        add_test_card(doc, number, category, color=TEAL)

    page_break(doc)
    add_page_heading(doc, "agents", "Korjaus ja puolustus", "Tee yksi rajattu korjaus, aja sama testi uudelleen ja valmistele 2–3 minuutin puolustus.", sublabel="OSA 3 · AGENTIT · VIIMEISTELY")
    add_h2(doc, "Yksi korjaus ja uudelleentesti", TEAL)
    add_label_detail_table(doc, [("Alkuperäinen tulos", "Kirjoita tähän."), ("Tehty korjaus", "Kirjoita tähän."), ("Saman testin uusi tulos", "Kirjoita tähän."), ("Johtopäätös", "Kirjoita tähän.")], color=TEAL, row_height=9)
    add_h2(doc, "Turvallisuusraja ja rehellinen rajoitus", TEAL)
    add_field(doc, "Mikä oikeus, hyväksyntä tai havaittava eskalointiehto rajaa toimintaa? Mitä työsi ei vielä todista?", height=34, color=TEAL)
    add_h2(doc, "Puolustus, 2–3 minuuttia", TEAL)
    add_field(doc, "Ongelma ja polku · aito rajattu mallivalinta · testi ja korjaus · turvallisuusraja ja rajoitus", height=38, color=TEAL)
    add_h2(doc, "Lopputarkistus", TEAL)
    add_checklist(doc, ["Työnkulussa on vähintään yksi aito rajattu kielimallin tekemä valinta.", "Kuusi rakennusosaa on käsitelty kanonisilla nimillä.", "Normaali, reuna- ja turvallisuustesti on dokumentoitu.", "Yksi korjaus ja sama uudelleentesti näkyvät ennen–jälkeen-vertailuna.", "Turvallisuusraja, rajoitus ja puolustus osoittavat, että ymmärrän kokonaisuuden."], color=TEAL, fill=TEAL_TINT)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, "KOKO KURSSI")
    add_kicker(doc, "KOKO KURSSI", KICKER)
    add_h1(doc, "Lähteet", BLUE)
    add_body(doc, "Kirjaa lähteet silloin, kun käytät niitä. Merkitse myös, mitä tarkistit ja milloin tarkistit sen.")
    table = doc.add_table(rows=9, cols=3)
    widths = [3900, CONTENT_W_DXA - 3900 - 1900, 1900]
    apply_table_geometry(table, widths)
    set_table_borders(table, BORDER, 5, inside=True, left_color=BLUE, left_size=10)
    for i, header in enumerate(("Lähde tai linkki", "Mitä tarkistin?", "Tarkistuspäivä")):
        shade_cell(table.rows[0].cells[i], BLUE_TINT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), FONT_BODY, 9, DARK, bold=True)
    for row in table.rows[1:]:
        row.height = Mm(18)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run("Kirjoita tähän."), FONT_BODY, 8.5, KICKER, italic=True)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    add_info_box(doc, "Säilytä oma työsi", "Tallenna valmis työkirja itsellesi Word- ja PDF-muodossa. Poista tai anonymisoi tiedot, joita et kurssin jälkeen enää tarvitse.", color=BLUE, fill=BLUE_TINT)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = "AI · Perusteet — työkirja"
    core.subject = "Kokeet, havainnot ja lopputyöt"
    core.author = "AI · Perusteet"
    core.keywords = "tekoäly, työkirja, prompti, botti, agentti"
    core.comments = "Kurssimateriaali: aiperusteet.fi · CC BY-SA 4.0"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Kirjoitettu {path} ({path.stat().st_size} tavua)")
